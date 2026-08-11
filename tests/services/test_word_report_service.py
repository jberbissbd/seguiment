from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
import pytest

from tutopy.application import create_services
from tutopy.models.messaging import CategoryNew, NoteNew, StudentNew
from tutopy.services.exceptions import EntityNotFoundError, ValidationError


def _scenario(db, tmp_path):
    services = create_services(db)
    student = services.students.create(StudentNew("Laia", "Martí", "4t A"))
    academic = services.categories.create(CategoryNew("Acadèmic"))
    family = services.categories.create(CategoryNew("Família"))
    unused = services.categories.create(CategoryNew("Sense notes"))
    services.report_configuration.set_category_order([family.id, academic.id, unused.id])
    services.notes.create(NoteNew(
        student.id, academic.id, "2026-02-01", 0, "Bona evolució"
    ))
    services.notes.create(NoteNew(
        student.id, family.id, "2025-10-10", 0, "Entrevista inicial"
    ))
    services.notes.create(NoteNew(
        student.id, academic.id, "2024-11-03", 0, "Curs anterior"
    ))
    return services, student, tmp_path / "informe"


def test_exporta_cursos_categories_ordenades_i_taules(db, tmp_path):
    services, student, destination = _scenario(db, tmp_path)
    path = services.word_reports.export_student(student.id, destination)
    assert path.suffix == ".docx"
    document = Document(path)
    headings = [
        (paragraph.text, paragraph.style.name)
        for paragraph in document.paragraphs
        if paragraph.style.type == WD_STYLE_TYPE.PARAGRAPH
        and paragraph.style.name.startswith("Heading")
    ]
    assert headings == [
        ("2024-2025", "Heading 1"),
        ("Acadèmic", "Heading 2"),
        ("2025-2026", "Heading 1"),
        ("Família", "Heading 2"),
        ("Acadèmic", "Heading 2"),
    ]
    assert [[cell.text for cell in row.cells] for row in document.tables[0].rows] == [
        ["Data", "Anotació"], ["03/11/2024", "Curs anterior"]
    ]
    assert [[cell.text for cell in row.cells] for row in document.tables[1].rows] == [
        ["Data", "Anotació"], ["10/10/2025", "Entrevista inicial"]
    ]
    assert [[cell.text for cell in row.cells] for row in document.tables[2].rows] == [
        ["Data", "Anotació"], ["01/02/2026", "Bona evolució"]
    ]
    page_breaks = document.element.xpath('.//w:br[@w:type="page"]')
    assert len(page_breaks) == 1


def test_configura_bloc_inicial_a4_marges_i_taules_a_amplada_completa(db, tmp_path):
    services, student, destination = _scenario(db, tmp_path)
    document = Document(services.word_reports.export_student(student.id, destination))
    section = document.sections[0]
    assert section.header.paragraphs[0].text == ""
    assert document.paragraphs[0].text == "Laia Martí"
    assert document.paragraphs[0].style.name == "Title"
    assert document.paragraphs[1].text == "Grup: 4t A"
    assert document.paragraphs[1].runs[0].bold
    assert section.page_width == pytest.approx(Cm(21), abs=1_000)
    assert section.page_height == pytest.approx(Cm(29.7), abs=1_000)
    assert section.top_margin == pytest.approx(Cm(2), abs=1_000)
    assert section.right_margin == pytest.approx(Cm(2), abs=1_000)
    assert section.bottom_margin == pytest.approx(Cm(2), abs=1_000)
    assert section.left_margin == pytest.approx(Cm(2), abs=1_000)
    for table in document.tables:
        table_width = table._tbl.tblPr.first_child_found_in("w:tblW")
        assert table_width.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "pct"
        assert table_width.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w") == "5000"
        assert table.autofit is False
        assert table.alignment == WD_TABLE_ALIGNMENT.CENTER
        assert table.columns[0].width == pytest.approx(Cm(3), abs=1_000)
        assert table.columns[1].width == pytest.approx(Cm(14), abs=1_000)


def test_permet_afegir_una_imatge_al_bloc_inicial(db, tmp_path):
    services, student, destination = _scenario(db, tmp_path)
    logo = tmp_path / "logo.png"
    logo.write_bytes(
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\rIDAT\x08\xd7c\xf8"
        b"\xcf\xc0\xf0\x1f\x00\x05\x00\x01\xff\x89\x99=\x1d\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    services.report_configuration.storage_dir = tmp_path / "reporting"
    services.report_configuration.set_header_image(logo)
    document = Document(services.word_reports.export_student(student.id, destination))
    assert len(document.inline_shapes) == 1


def test_rebutja_una_imatge_de_capcalera_inexistent(db, tmp_path):
    services, _student, _destination = _scenario(db, tmp_path)
    services.report_configuration.storage_dir = tmp_path / "reporting"
    with pytest.raises(ValidationError, match="imatge de capçalera"):
        services.report_configuration.set_header_image(tmp_path / "inexistent.png")


def test_rebutja_alumne_inexistent_o_sense_notes(db, tmp_path):
    services = create_services(db)
    with pytest.raises(EntityNotFoundError):
        services.word_reports.export_student(999, tmp_path / "x.docx")
    student = services.students.create(StudentNew("Pau", "Puig", "3r A"))
    with pytest.raises(ValidationError, match="no té notes"):
        services.word_reports.export_student(student.id, tmp_path / "x.docx")


def test_elimina_controls_xml_no_valids(db, tmp_path):
    services = create_services(db)
    student = services.students.create(StudentNew("Júlia", "Nuñez", "4t A"))
    category = services.categories.create(CategoryNew("Acadèmic"))
    services.notes.create(NoteNew(
        student.id, category.id, "2026-02-01", 0, "Evolució\x07 positiva 😊"
    ))
    path = services.word_reports.export_student(student.id, tmp_path / "unicode.docx")
    assert Document(path).tables[0].cell(1, 1).text == "Evolució positiva 😊"
