"""Generació de l'informe DOCX de les notes de seguiment d'un alumne."""

from collections import defaultdict
from datetime import date
from pathlib import Path

from tutopy.database.daos.academic_course_dao import AcademicCourseDAO
from tutopy.database.daos.note_dao import NoteDAO
from tutopy.database.daos.student_dao import StudentDAO
from tutopy.services.exceptions import EntityNotFoundError, ValidationError
from tutopy.services.report_configuration_service import ReportConfigurationService
from tutopy.services.validation_service import ValidationService
from tutopy.services.utils import sanitize_xml_text


class WordReportService:
    """Genera un informe DOCX de les notes de seguiment d'un alumne."""

    def __init__(self, students: StudentDAO, notes: NoteDAO,
                 courses: AcademicCourseDAO,
                 configuration: ReportConfigurationService, batch_loader=None):
        """Rep els repositoris de domini i, opcionalment, el carregador de lots."""
        self.students = students
        self.notes = notes
        self.courses = courses
        self.configuration = configuration
        self.batch_loader = batch_loader
        self.validation = ValidationService()

    def export_student(self, student_id: int, destination: str | Path) -> Path:
        """Exporta l'informe DOCX d'un únic alumne.

        Args:
            student_id: ID de l'alumne.
            destination: Ruta de destinació (l'extensió es normalitza a `.docx`).

        Returns:
            Ruta final del fitxer generat.

        Raises:
            ValidationError: Si l'alumne no té notes per exportar.
            EntityNotFoundError: Si l'alumne no existeix.
        """
        student_id = self.validation.positive_id(student_id)
        data = self.prepare_students([student_id])
        return self.export_prepared(student_id, destination, data)

    def prepare_students(self, student_ids: list[int]):
        """Carrega una vegada les dades compartides dels informes DOCX."""
        return self.batch_loader.load(student_ids, include_header=True)

    def export_prepared(self, student_id: int, destination: str | Path, data) -> Path:
        """Genera un DOCX a partir d'un snapshot carregat prèviament."""
        student, student_notes = self._student_and_notes(student_id, data)
        path = self._document_path(destination)
        document = self._new_document(student, data.header_image)
        by_course = self._group_by_course(student_notes)
        self._add_courses(
            document, by_course, data.categories,
            self._course_names(by_course, data.course_names),
        )
        self._save_document(document, path)
        return path

    @staticmethod
    def _student_and_notes(student_id: int, data):
        student = data.students.get(student_id)
        if student is None:
            raise EntityNotFoundError("L’alumne seleccionat no existeix.")
        student_notes = sorted(
            data.notes[student_id], key=lambda note: (note.date, note.id)
        )
        if not student_notes:
            raise ValidationError("L’alumne no té notes per exportar.")
        return student, student_notes

    @staticmethod
    def _document_path(destination: str | Path) -> Path:
        path = Path(destination)
        if path.suffix.lower() != ".docx":
            path = path.with_suffix(".docx")
        if not path.name:
            raise ValidationError("Cal indicar una destinació per a l’informe.")
        return path

    @staticmethod
    def _group_by_course(student_notes):
        by_course = defaultdict(list)
        for note in student_notes:
            by_course[note.course_id].append(note)
        return by_course

    def _new_document(self, student, header_image):
        from docx import Document

        document = Document()
        document.core_properties.title = self._safe_text(
            f"Informe de {student.filing_name}"
        )
        document.core_properties.subject = "Notes de seguiment"
        self._configure_page(document)
        self._add_student_header(
            document, student, header_image
        )
        return document

    def _add_courses(self, document, by_course, categories, course_names) -> None:
        for index, (course_id, course_notes) in enumerate(sorted(
            by_course.items(), key=lambda item: course_names[item[0]]
        )):
            if index:
                document.add_page_break()
            self._add_course(
                document, course_names[course_id], course_notes, categories
            )

    def _add_course(self, document, course_name, course_notes, categories) -> None:
        document.add_heading(course_name, level=1)
        notes_by_category = defaultdict(list)
        for note in course_notes:
            notes_by_category[note.category_id].append(note)
        for category in categories:
            category_notes = notes_by_category.get(category.id)
            if category_notes:
                self._add_category(document, category.name, category_notes)

    def _add_category(self, document, category_name: str, category_notes) -> None:
        document.add_heading(self._safe_text(category_name), level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        self._configure_table_width(table)
        self._configure_table_header(table.rows[0].cells)
        for note in category_notes:
            self._add_note_row(table, note)

    @staticmethod
    def _configure_table_header(headers) -> None:
        headers[0].text = "Data"
        headers[1].text = "Anotació"
        for cell in headers:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    def _add_note_row(self, table, note) -> None:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Cm

        cells = table.add_row().cells
        cells[0].text = date.fromisoformat(note.date).strftime("%d/%m/%Y")
        cells[1].text = self._safe_text(note.content)
        for cell, width in zip(cells, (Cm(3), Cm(14)), strict=True):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    @staticmethod
    def _save_document(document, path: Path) -> None:
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            document.save(path)
        except (OSError, ValueError) as error:
            raise ValidationError("No s’ha pogut desar l’informe.") from error

    @staticmethod
    def _configure_page(document) -> None:
        """Configura un A4 amb marges de 2 cm."""
        from docx.shared import Cm

        section = document.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.right_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)

    def _add_student_header(self, document, student, header_image) -> None:
        """Afegeix el bloc inicial identificatiu i, opcionalment, un logotip."""
        from docx.image.exceptions import UnrecognizedImageError
        from docx.shared import Cm

        if header_image:
            image_path = Path(header_image)
            if not image_path.is_file():
                raise ValidationError("No s’ha trobat la imatge de capçalera.")
            try:
                document.add_picture(str(image_path), width=Cm(5))
            except (OSError, ValueError, UnrecognizedImageError) as error:
                raise ValidationError("La imatge de capçalera no és vàlida.") from error
        document.add_heading(self._safe_text(student.filing_name), level=0)
        group = self._safe_text(student.group_name) or "Sense grup"
        paragraph = document.add_paragraph()
        paragraph.add_run("Grup: ").bold = True
        paragraph.add_run(group)
        exported = document.add_paragraph()
        exported.add_run("Data d’exportació: ").bold = True
        exported.add_run(date.today().strftime("%d/%m/%Y"))

    @staticmethod
    def _configure_table_width(table) -> None:
        """Fixa la taula al 100% dels 17 cm disponibles entre marges."""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm

        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(14)
        table_width = table._tbl.tblPr.first_child_found_in("w:tblW")
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            table._tbl.tblPr.append(table_width)
        table_width.set(qn("w:type"), "pct")
        table_width.set(qn("w:w"), "5000")
        for cell, width in zip(table.rows[0].cells, (Cm(3), Cm(14)), strict=True):
            cell.width = width

    @staticmethod
    def _course_names(by_course, names) -> dict[int, str]:
        missing = set(by_course) - names.keys()
        if missing:
            raise ValidationError("Una nota fa referència a un curs acadèmic inexistent.")
        return names

    @staticmethod
    def _safe_text(value) -> str:
        """Elimina els controls que XML no admet, conservant salts i tabuladors."""
        return sanitize_xml_text(value)
